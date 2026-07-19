"""
doc_reader.py
Modul untuk membaca file Word (.doc dan .docx) dengan fallback berlapis.

Urutan strategi untuk .doc:
1. Convert via LibreOffice (soffice) -> .docx -> baca pakai python-docx
2. Fallback: antiword (kalau LibreOffice tidak tersedia/gagal)
3. Kalau semua gagal, raise error yang jelas

Untuk .docx: langsung pakai python-docx.
"""

import os
import subprocess
import tempfile
import shutil
import logging

logger = logging.getLogger(__name__)

try:
    import docx  # python-docx
except ImportError:
    docx = None


class DocReadError(Exception):
    """Raised ketika semua metode ekstraksi gagal."""
    pass


def _check_tool_available(tool_name: str) -> bool:
    """Cek apakah command line tool tersedia di PATH."""
    return shutil.which(tool_name) is not None


def _convert_doc_to_docx_libreoffice(doc_path: str, output_dir: str) -> str:
    """
    Convert .doc ke .docx menggunakan LibreOffice headless.
    Return path ke file .docx hasil konversi.
    """
    if not _check_tool_available("soffice"):
        raise DocReadError("LibreOffice (soffice) tidak ditemukan di PATH")

    try:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", output_dir,
                doc_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
    except subprocess.TimeoutExpired:
        raise DocReadError("Konversi LibreOffice timeout (>60 detik)")

    if result.returncode != 0:
        raise DocReadError(f"LibreOffice gagal convert: {result.stderr.strip()}")

    base_name = os.path.splitext(os.path.basename(doc_path))[0]
    converted_path = os.path.join(output_dir, f"{base_name}.docx")

    if not os.path.exists(converted_path):
        raise DocReadError(f"File hasil konversi tidak ditemukan: {converted_path}")

    return converted_path


def _extract_text_antiword(doc_path: str) -> str:
    """Fallback: ekstrak teks .doc langsung pakai antiword."""
    if not _check_tool_available("antiword"):
        raise DocReadError("antiword tidak ditemukan di PATH")

    try:
        result = subprocess.run(
            ["antiword", doc_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
    except subprocess.TimeoutExpired:
        raise DocReadError("antiword timeout (>30 detik)")

    if result.returncode != 0:
        raise DocReadError(f"antiword gagal: {result.stderr.strip()}")

    return result.stdout


def _extract_text_from_docx(docx_path: str) -> str:
    """Ekstrak teks dari file .docx menggunakan python-docx."""
    if docx is None:
        raise DocReadError(
            "python-docx belum terpasang. Jalankan: pip install python-docx"
        )

    try:
        document = docx.Document(docx_path)
    except Exception as e:
        raise DocReadError(f"Gagal membuka .docx: {e}")

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Ekstrak juga teks dari tabel, kalau ada
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def read_word_document(file_path: str) -> str:
    """
    Fungsi utama: baca file Word (.doc atau .docx) dan return teksnya.

    Args:
        file_path: path ke file .doc atau .docx

    Returns:
        str: teks yang berhasil diekstrak

    Raises:
        DocReadError: kalau semua metode gagal
        FileNotFoundError: kalau file tidak ada
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        logger.info("Membaca file .docx langsung", extra={"doc_path": file_path})
        return _extract_text_from_docx(file_path)

    elif ext == ".doc":
        logger.info("Memproses file .doc lama", extra={"doc_path": file_path})

        errors = []

        # Strategi 1: LibreOffice convert -> docx -> baca
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                converted_path = _convert_doc_to_docx_libreoffice(file_path, tmp_dir)
                text = _extract_text_from_docx(converted_path)
                logger.info("Berhasil via LibreOffice", extra={"doc_path": file_path})
                return text
        except DocReadError as e:
            errors.append(f"LibreOffice: {e}")
            logger.warning(
                "LibreOffice gagal, mencoba antiword",
                extra={"doc_path": file_path, "error_detail": str(e)},
            )

        # Strategi 2: fallback ke antiword
        try:
            text = _extract_text_antiword(file_path)
            logger.info("Berhasil via antiword", extra={"doc_path": file_path})
            return text
        except DocReadError as e:
            errors.append(f"antiword: {e}")

        # Semua gagal
        raise DocReadError(
            f"Gagal membaca file .doc dengan semua metode. Detail: {'; '.join(errors)}"
        )

    else:
        raise DocReadError(
            f"Ekstensi tidak didukung: {ext}. Hanya .doc dan .docx yang bisa diproses."
        )


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python doc_reader.py <path_ke_file.doc_atau_.docx>")
        sys.exit(1)

    logging.basicConfig(level=logging.INFO)

    try:
        content = read_word_document(sys.argv[1])
        print("=== TEKS BERHASIL DIEKSTRAK ===")
        print(content[:2000])
        print(f"\n... (total {len(content)} karakter)")
    except (DocReadError, FileNotFoundError) as e:
        print(f"❌ Error: {e}")
        sys.exit(1)
