"""
CV Generator Agent — Specializes in rewriting and generating ATS-optimized CVs
using Harvard Business School resume standards.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import config
from llm_client import chat_completion

import json
from database import DatabaseManager
from sqlalchemy import text as sql_text


ATS_CV_PROMPT = """Kamu adalah CV Writer Expert profesional lulusan Harvard Business School dan spesialis optimasi ATS (Applicant Tracking System).
Tugasmu adalah menulis ulang CV pengguna menjadi format ATS yang sangat kuat, tajam, dan berdampak tinggi.

ATURAN WAJIB (STANDAR CV ATS-FRIENDLY):
1. **Struktur Konten yang Baku:** Gunakan format Markdown yang sangat rapi dan berurutan:
   - Contact Information
   - Professional Summary (3-4 kalimat padat)
   - Work Experience (Urutan kronologis terbalik, max 5-6 poin per posisi)
   - Education
   - Skills (Bagi menjadi Hard Skills & Soft Skills)
   - Certifications (Opsional)
   Gunakan header standar global di atas untuk kedua versi bahasa (atau padanannya yang standar).

2. **Action Verbs & Metrik:**
   Mulai SETIAP poin pengalaman kerja dengan action verb (kata kerja aktif).
   PERTAHANKAN dan SOROT angka/metrik kuantitatif yang SUDAH ADA di CV asli. Jangan hilangkan angka tersebut, tapi dilarang keras mengarang angka baru.

3. **Tanpa Keyword Stuffing:**
   Masukkan kata kunci dari lowongan kerja secara natural ke dalam Ringkasan, Pengalaman, dan Keahlian. Jangan sekadar me-list kata kunci tanpa konteks.

4. **Konsistensi Bahasa:**
   Kamu HARUS menghasilkan dua versi utuh: 1 versi Bahasa Indonesia penuh dan 1 versi Bahasa Inggris penuh. Jangan campur bahasa dalam satu versi.

5. **GUARDRAIL EKSPLISIT (SANGAT KRUSIAL):**
   Jangan pernah menambahkan pengalaman kerja, gelar, sertifikasi, atau pencapaian yang tidak disebutkan di CV asli user. Tugas kamu adalah menata ulang dan memperkuat bahasa dari apa yang sudah ada, bukan mengarang yang baru.

6. **Formula XYZ Google (WAJIB di semua poin pengalaman kerja):**
   Tulis setiap poin pencapaian menggunakan pola:
   "[Action Verb] [X — apa yang dilakukan/dibuat] dengan [Y — cara/metode] sehingga menghasilkan [Z — dampak kuantitatif/kualitatif]."
   Contoh kuat: "Memimpin tim 5 orang dalam merancang sistem backend berbasis FastAPI, berhasil memangkas waktu respons API sebesar 40% dalam 2 sprint."
   Contoh lemah (HINDARI): "Bertanggung jawab atas pengembangan backend."

{rubric_context}

Outputmu harus MURNI dan HANYA sebuah JSON terstruktur tanpa markdown formatting block, dengan skema berikut:
{
  "cv_indonesia": "CV lengkap versi Bahasa Indonesia dalam format Markdown...",
  "cv_english": "CV lengkap versi Bahasa Inggris dalam format Markdown..."
}
"""

def get_scoring_rubric_context() -> str:
    """Fetch scoring rubric from database to inject into prompt."""
    try:
        db = DatabaseManager()
        with db.engine.connect() as conn:
            result = conn.execute(sql_text("SELECT kriteria FROM scoring_rubric"))
            criteria = [row[0] for row in result]
            if criteria:
                rubric_text = "\n- ".join(criteria)
                return f"\nKRITERIA SCORING RUBRIC (Pastikan CV memenuhi ini):\n- {rubric_text}\n"
    except Exception as e:
        print(f"Warning: Failed to fetch scoring rubric: {e}")
    return ""


def generate_ats_cv(cv_text: str, target_job: dict = None, language: str = "auto") -> dict:
    """
    Generate an ATS-friendly version of the CV using Harvard standards.
    Generates both Indonesian and English versions in one call.
    """
    if not config.is_llm_configured():
        return {"ats_text_id": None, "ats_text_en": None, "available": False}

    try:
        rubric_context = get_scoring_rubric_context()
        system_prompt = ATS_CV_PROMPT.replace("{rubric_context}", rubric_context)
        
        target_context = ""
        if target_job:
            target_context = f"\n\nPosisi Target:\n- Jabatan: {target_job.get('job_title', 'N/A')}\n- Perusahaan: {target_job.get('company_name', 'N/A')}\n- Deskripsi Pekerjaan: {target_job.get('job_description', 'N/A')}"
            system_prompt += f"\n\nPENTING: Optimalkan CV ini KHUSUS untuk posisi target di bawah secara natural."

        reply = chat_completion(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"CV Asli/Data Mentah:\n\n{cv_text}{target_context}"},
            ],
            temperature=0.3, # lower temp to reduce hallucinations and ensure valid JSON
            max_tokens=6000,
            agent_id=2,
        )
        
        # Parse JSON output
        try:
            # Clean possible markdown block if LLM added it despite instructions
            clean_reply = reply.strip()
            if clean_reply.startswith("```json"):
                clean_reply = clean_reply[7:-3].strip()
            elif clean_reply.startswith("```"):
                clean_reply = clean_reply[3:-3].strip()
                
            parsed = json.loads(clean_reply)
            cv_id = parsed.get("cv_indonesia", "Gagal parse Bahasa Indonesia.")
            cv_en = parsed.get("cv_english", "Gagal parse English.")
            
            return {
                "ats_text_id": cv_id,
                "ats_text_en": cv_en,
                "available": True,
            }
        except json.JSONDecodeError as e:
            return {"ats_text_id": f"Error parsing JSON dari LLM: {e}\n\nRaw output:\n{reply}", "ats_text_en": None, "available": True}

    except Exception as e:
        return {"ats_text_id": f"Error: {str(e)}", "ats_text_en": None, "available": True}


def export_cv_to_docx(cv_text: str) -> bytes:
    """Export CV text to a formatted DOCX file. Returns bytes."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)

    # Parse and add content
    lines = cv_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Detect headings (lines in ALL CAPS or starting with ##)
        if line.startswith("## ") or line.startswith("# "):
            clean = line.lstrip("#").strip()
            p = doc.add_heading(clean, level=2)
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            p = doc.add_heading(line, level=2)
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            p = doc.add_paragraph(clean, style="List Bullet")
        else:
            p = doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_cv_to_pdf(cv_text: str) -> bytes:
    """Export CV text to a formatted PDF file. Returns bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=40,
        bottomMargin=40,
    )
    styles = getSampleStyleSheet()

    # Custom styles
    h1_style = ParagraphStyle(
        "Heading1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        spaceAfter=12,
        textColor="#1f2937",
    )
    h2_style = ParagraphStyle(
        "Heading2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        spaceAfter=8,
        spaceBefore=12,
        textColor="#374151",
    )
    normal_style = ParagraphStyle(
        "Normal",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        spaceAfter=6,
        leading=14,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=normal_style,
        leftIndent=20,
        firstLineIndent=-10,
    )

    story = []
    lines = cv_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue

        if line.startswith("# "):
            clean = line.lstrip("#").strip()
            story.append(Paragraph(clean, h1_style))
        elif line.startswith("## "):
            clean = line.lstrip("#").strip()
            story.append(Paragraph(clean, h2_style))
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            story.append(Paragraph(line, h2_style))
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            story.append(Paragraph(f"• {clean}", bullet_style))
        else:
            story.append(Paragraph(line, normal_style))

    doc.build(story)
    return buffer.getvalue()
