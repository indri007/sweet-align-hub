"""
CV Processor (Document) — Extract text from PDF and DOCX files.
Terintegrasi ke dalam arsitektur `core/` baru.
"""

import io
from pathlib import Path

def extract_text_from_pdf(file_bytes: bytes, gemini_client=None) -> str:
    """Extract text content from a PDF file with Gemini Vision OCR fallback for scanned PDFs."""
    import pdfplumber
    text_parts = []
    
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
                
    text = "\\n\\n".join(text_parts).strip()
    
    # OCR Fallback if text is empty or too short (likely scanned PDF)
    if len(text) < 50:
        if gemini_client:
            try:
                from google.genai import types
                response = gemini_client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[
                        types.Part.from_bytes(data=file_bytes, mime_type="application/pdf"),
                        "Ekstrak dan ketik ulang seluruh teks yang ada di dalam dokumen CV ini secara lengkap, pertahankan urutan dan bahasanya.",
                    ]
                )
                if response.text:
                    return response.text.strip()
            except Exception as e:
                print(f"[Document/OCR] Gemini Vision failed: {e}")
                
        # Fallback 2: Local Tesseract OCR via pdf2image + pytesseract
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            images = convert_from_bytes(file_bytes)
            ocr_parts = []
            for image in images:
                ocr_text = pytesseract.image_to_string(image)
                if ocr_text:
                    ocr_parts.append(ocr_text)
            ocr_result = "\\n\\n".join(ocr_parts).strip()
            if ocr_result:
                return ocr_result
        except Exception as e:
            print(f"[Document/OCR] Local Tesseract failed: {e}")

    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
            
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\\n".join(text_parts)

def extract_cv_text(file_bytes: bytes, filename: str, gemini_client=None) -> str:
    """
    Extract text from a CV file based on its extension.
    Supports PDF and DOCX formats.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes, gemini_client)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".doc":
        raise ValueError(
            "Format .doc (Word lama/binary) tidak didukung. "
            "Silakan simpan ulang CV kamu sebagai .docx atau PDF, lalu upload lagi."
        )
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")

def validate_cv_file(file_bytes: bytes, filename: str, max_size_mb: int = 100) -> tuple[bool, str]:
    """
    Validate CV file format and size.
    Returns (is_valid, error_message).
    """
    ext = Path(filename).suffix.lower()
    if ext == ".doc":
        return False, "Format .doc (Word lama) tidak didukung. Simpan ulang sebagai .docx atau PDF."
    if ext not in (".pdf", ".docx"):
        return False, f"Format {ext} tidak didukung. Gunakan PDF atau DOCX."

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return False, f"Ukuran file {size_mb:.1f}MB melebihi batas {max_size_mb}MB."

    return True, ""
