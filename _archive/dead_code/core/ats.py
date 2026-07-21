from dataclasses import dataclass, field
from typing import Optional
from .models import CandidateProfile, JobPosting, MatchCategory
from .pipeline import EmbeddingPipeline
import re
import json
import io

# === PROMPTS ASLI ===
REVIEW_PROMPT = """You are a professional ATS (Applicant Tracking System) consultant with over 40 years of experience in recruitment and HR tech. Your task is to analyze the following CV and provide an objective assessment of its ATS-compatibility, based on these criteria:

1. FORMAT & STRUCTURE (weight 25%)
   - Does it use a single-column layout (not multi-column)?
   - Is it free of text boxes, complex tables, or graphic elements that disrupt parsing?
   - Do section headings use standard terminology (Work Experience, Education, Skills)?
   - Is the font machine-readable (Arial, Calibri, Helvetica, Georgia)?

2. KEYWORDS & RELEVANCE (weight 30%)
   - Does the CV contain keywords relevant to the target position?
   - Are technical terms/skills written using industry-standard terminology (not uncommon abbreviations)?

3. QUANTIFIED RESULTS (weight 20%)
   - Is work experience accompanied by measurable metrics (percentages, figures, counts)?
   - Does it use active verbs (led, developed, increased) rather than passive voice?

4. COMPLETENESS & CONSISTENCY (weight 15%)
   - Is date formatting consistent throughout the document?
   - Is contact information (email, phone, city, LinkedIn) complete and easy to locate?

5. PROFESSIONAL VISUAL IMPRESSION (weight 10%)
   - Are there visual elements (horizontal divider lines, spacing) that make the CV easy for human eyes to scan, without compromising ATS parsing?

OUTPUT INSTRUCTIONS — provide results EXACTLY in this format (use these exact headings):

## 📊 ATS Score: [score]/100

## 📋 Skor Per Kategori
- Format & Struktur: [x]/25
- Kata Kunci & Relevansi: [x]/30
- Kuantifikasi Hasil: [x]/20
- Kelengkapan & Konsistensi: [x]/15
- Kesan Profesional Visual: [x]/10

## ✅ Kekuatan Utama (3-5 poin)
- [specific finding, NOT generic praise]

## ⚠️ Area yang Perlu Diperbaiki (3-5 poin)
- [concrete and actionable suggestion, NOT vague advice]

## 🔑 Keywords yang Terdeteksi
[list detected keywords/skills found in the CV]

## 📝 Ringkasan Kesiapan ATS
[One paragraph maximum 3 sentences summarizing how ready this CV is to pass ATS filters. No generic praise.]

Respond in Bahasa Indonesia. Focus on SPECIFIC, actionable findings only."""

ATS_CV_PROMPT = """Kamu adalah Perekrut Senior dan Pakar CV ATS (Applicant Tracking System). Tugasmu adalah menulis ulang CV asli user menjadi versi CV ATS-friendly yang setara dengan standar "Harvard Resume Format".

Aturan Mutlak Penulisan (Berdasarkan Knowledge Base HRD):
1. **Format Single-Column:** Jangan gunakan kolom ganda, tabel rumit, atau elemen grafis. Gunakan pemisah garis lurus jika perlu.
2. **Standard Headings:** Wajib gunakan urutan section baku: [PROFIL / PROFESSIONAL SUMMARY], [PENGALAMAN KERJA / WORK EXPERIENCE], [PENDIDIKAN / EDUCATION], [SKILLS & TOOLS], [SERTIFIKASI / CERTIFICATIONS].
3. **Formula Bullet Points (XYZ Method):** Ubah deskripsi pekerjaan user menjadi format "Telah mencapai [X] yang diukur dengan [Y], dengan melakukan [Z]".
4. **Action Verbs yang Kuat:** Awali setiap bullet point pengalaman kerja dengan kata kerja aktif yang kuat (misal: Mengembangkan, Memimpin, Meningkatkan, Merancang, Mengoptimalkan).
5. **Kuantifikasi:** Tambahkan metrik/angka fiktif logis JIKA user tidak menuliskannya, namun beri tanda kurung siku (misal: [meningkatkan efisiensi 20%]) agar user tahu mereka harus mengisinya dengan angka asli mereka.
6. **Optimasi Keyword:** Ekstrak dan tonjolkan Hard Skills yang ada di CV asli ke dalam list yang mudah dipindai oleh bot ATS.

Output HANYA isi teks CV yang sudah diperbaiki secara menyeluruh, siap di-copy-paste oleh user. JANGAN berikan kalimat pembuka/penutup seperti "Berikut adalah CV Anda".
Gunakan format Markdown untuk struktur Heading dan Bullet Points."""

@dataclass
class ATSScoreResult:
    total_score: float                 
    keyword_match_score: float
    format_compliance_score: float
    semantic_similarity_score: float
    breakdown_by_category: dict        
    missing_items: list[str]           
    weights: dict = field(default_factory=lambda: {
        "keyword_match": 0.4,
        "format_compliance": 0.2,
        "semantic_similarity": 0.4,
    })


class ATSScorer:
    """Hitung ATS score + persentase match kandidat vs job spesifik."""
    def __init__(self, embedding_pipeline: EmbeddingPipeline):
        self.embedding_pipeline = embedding_pipeline

    def score_keyword_match(self, cv: CandidateProfile, job: JobPosting) -> float:
        job_keywords = set(re.findall(r'\b[a-zA-Z]{3,}\b', job.requirements.lower()))
        if not job_keywords: return 100.0
        cv_text = cv.raw_cv_text.lower()
        matched = sum(1 for kw in job_keywords if kw in cv_text)
        return (matched / len(job_keywords)) * 100.0

    def score_format_compliance(self, raw_cv_text: str) -> float:
        score = 100.0
        if len(raw_cv_text) < 500: score -= 50
        lower_text = raw_cv_text.lower()
        if "experience" not in lower_text and "pengalaman" not in lower_text: score -= 20
        if "education" not in lower_text and "pendidikan" not in lower_text: score -= 20
        return max(0.0, score)

    def score_semantic_similarity(self, cv: CandidateProfile, job: JobPosting) -> float:
        cv_vec = self.embedding_pipeline.embed_text(cv.raw_cv_text)
        job_vec = self.embedding_pipeline.embed_text(f"{job.title} {job.description}")
        import math
        dot_product = sum(a*b for a, b in zip(cv_vec, job_vec))
        norm_a = math.sqrt(sum(a*a for a in cv_vec))
        norm_b = math.sqrt(sum(b*b for b in job_vec))
        if norm_a == 0 or norm_b == 0: return 0.0
        similarity = dot_product / (norm_a * norm_b)
        return max(0.0, min(100.0, (similarity + 1) / 2 * 100))

    def _fetch_weights(self) -> dict:
        from database import DatabaseManager
        from sqlalchemy import text
        weights = {
            "keyword_match": 0.4,
            "format_compliance": 0.2,
            "semantic_similarity": 0.4,
        }
        try:
            db = DatabaseManager()
            with db.engine.connect() as conn:
                result = conn.execute(text("SELECT dimension, weight FROM scoring_rubric"))
                for row in result:
                    # SQLAlchemy row[0] = dimension, row[1] = weight
                    dim = row[0]
                    if dim in weights:
                        weights[dim] = float(row[1])
        except Exception as e:
            print(f"Warning: Failed to fetch weights from Aiven: {e}")
        return weights

    def compute(self, cv: CandidateProfile, job: JobPosting) -> ATSScoreResult:
        weights = self._fetch_weights()
        kw_score = self.score_keyword_match(cv, job)
        fmt_score = self.score_format_compliance(cv.raw_cv_text)
        sem_score = self.score_semantic_similarity(cv, job)
        total = (kw_score * weights['keyword_match']) + (fmt_score * weights['format_compliance']) + (sem_score * weights['semantic_similarity'])
        breakdown = {MatchCategory.SKILL: kw_score, MatchCategory.EXPERIENCE: sem_score, MatchCategory.CULTURE: sem_score * 0.8}
        return ATSScoreResult(
            total_score=round(total, 2), keyword_match_score=round(kw_score, 2),
            format_compliance_score=round(fmt_score, 2), semantic_similarity_score=round(sem_score, 2),
            breakdown_by_category=breakdown, missing_items=[], weights=weights
        )

def detect_cv_language(cv_text: str) -> str:
    """
    Deteksi otomatis bahasa CV: 'id' (Indonesia) atau 'en' (English).
    Menghitung kemunculan kata umum khas masing-masing bahasa.
    """
    text_lower = cv_text.lower()

    indonesian_markers = [
        "pengalaman", "pendidikan", "keterampilan", "keahlian",
        "riwayat", "pekerjaan", "universitas", "sekolah",
        "tanggung jawab", "prestasi", "kemampuan",
    ]
    english_markers = [
        "experience", "education", "skills", "employment",
        "responsibilities", "achievements", "university",
        "objective", "summary", "certifications",
    ]

    id_score = sum(text_lower.count(word) for word in indonesian_markers)
    en_score = sum(text_lower.count(word) for word in english_markers)

    return "id" if id_score >= en_score else "en"


def resolve_output_language(cv_text: str, output_language: str = "auto") -> str:
    """
    Menentukan bahasa output akhir.
    output_language: 'auto', 'id', atau 'en'
    """
    if output_language in ("id", "en"):
        return output_language
    return detect_cv_language(cv_text)

class CVGenerator:
    """Generate/rewrite CV grounded di data user — TIDAK BOLEH hallucinate."""
    def __init__(self, gemini_client):
        self.gemini_client = gemini_client

    def generate_id(self, candidate: CandidateProfile, target_job: Optional[JobPosting] = None) -> str:
        prompt = ATS_CV_PROMPT + "\n\nOutput CV ATS-friendly HARUS ditulis dalam Bahasa Indonesia.\n"
        if target_job: prompt += f"Sesuaikan highlight dengan lowongan: {target_job.title}. JANGAN menambahkan fakta fiktif.\n"
        prompt += f"\nCV Asli:\n{candidate.raw_cv_text}"
        response = self.gemini_client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        return response.text

    def generate(self, candidate: CandidateProfile, output_language: str = "auto", target_job: Optional[JobPosting] = None) -> str:
        lang = resolve_output_language(candidate.raw_cv_text, output_language)
        if lang == "id":
            return self.generate_id(candidate, target_job)
        return self.generate_en(candidate, target_job)

    def generate_en(self, candidate: CandidateProfile, target_job: Optional[JobPosting] = None) -> str:
        prompt = ATS_CV_PROMPT + "\n\nOutput CV ATS-friendly HARUS ditulis dalam Bahasa Inggris (English).\n"
        if target_job: prompt += f"Tailor the summary and highlights for this job: {target_job.title}. DO NOT hallucinate facts.\n"
        prompt += f"\nOriginal CV:\n{candidate.raw_cv_text}"
        response = self.gemini_client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        return response.text

    def validate_no_hallucination(self, generated_cv: str, source: CandidateProfile) -> bool:
        prompt = f"Apakah teks CV Baru mengandung fakta, skill, atau pengalaman kerja fiktif yang TIDAK ADA di CV Asli? Jawab dengan JSON {{'has_hallucination': true/false}}.\n\nCV Asli:\n{source.raw_cv_text}\n\nCV Baru:\n{generated_cv}"
        response = self.gemini_client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        try:
            text = response.text.replace("```json", "").replace("```", "").strip()
            data = json.loads(text)
            return not data.get("has_hallucination", False)
        except Exception:
            return True

def _add_bottom_border_to_docx_paragraph(paragraph):
    """Menambahkan garis bawah profesional (bottom border) pada Heading DOCX."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5 pt
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '000000') # Solid black
    pBdr.append(bottom)
    pPr.append(pBdr)

def export_cv_to_docx(cv_text: str) -> bytes:
    """Export CV text to a formatted DOCX file dengan garis profesional ATS."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)

    lines = cv_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("## ") or line.startswith("# "):
            clean = line.lstrip("#").strip()
            p = doc.add_heading(clean, level=2)
            _add_bottom_border_to_docx_paragraph(p)
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            p = doc.add_heading(line, level=2)
            _add_bottom_border_to_docx_paragraph(p)
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            p = doc.add_paragraph(clean, style="List Bullet")
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def export_cv_to_pdf(cv_text: str) -> bytes:
    """Export CV text to a formatted PDF file dengan garis profesional ATS."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.platypus.flowables import HRFlowable

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle(
        "CVHeading", parent=styles["Heading2"], fontSize=13, spaceAfter=2, spaceBefore=14, textColor="#1a1a2e"
    )
    body_style = ParagraphStyle("CVBody", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14)
    bullet_style = ParagraphStyle(
        "CVBullet", parent=styles["Normal"], fontSize=10, spaceAfter=3, leftIndent=20, leading=14, bulletIndent=10
    )

    elements = []
    lines = cv_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 6))
            continue
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if line.startswith("## ") or line.startswith("# "):
            clean = safe_line.lstrip("#").strip()
            elements.append(Paragraph(clean, heading_style))
            elements.append(HRFlowable(width="100%", thickness=1, color="black", spaceBefore=2, spaceAfter=6))
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            elements.append(Paragraph(safe_line, heading_style))
            elements.append(HRFlowable(width="100%", thickness=1, color="black", spaceBefore=2, spaceAfter=6))
        elif line.startswith("- ") or line.startswith("• "):
            clean = safe_line.lstrip("-•").strip()
            elements.append(Paragraph(f"• {clean}", bullet_style))
        else:
            elements.append(Paragraph(safe_line, body_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
