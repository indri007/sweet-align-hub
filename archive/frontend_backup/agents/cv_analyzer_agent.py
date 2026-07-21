"""
CV Analyzer Agent — Reviews CV content and generates feedback.
Provides ATS score, improvement suggestions, and generates ATS-friendly CV.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import config
from llm_client import chat_completion


REVIEW_PROMPT = """You are a professional ATS (Applicant Tracking System) consultant with over 40 years of experience in recruitment and HR tech. Your task is to analyze the following CV and provide an objective assessment of its ATS-compatibility, based on these criteria:

1. FORMAT & STRUCTURE (weight 25%)
   - Does it use a single-column layout (not multi-column)?
   - Is it free of text boxes, complex tables, or graphic elements that disrupt parsing?
   - Do section headings use standard terminology (Work Experience, Education, Skills)?
   - Is the font machine-readable (Arial, Calibri, Helvetica, Georgia)?

2. KEYWORDS & RELEVANCE (weight 30%)
   - Does the CV contain keywords relevant to the target position?
   - Are technical terms/skills written using industry-standard terminology (not uncommon abbreviations)?

3. QUANTIFIED RESULTS (weight 20%)
   - Is work experience accompanied by measurable metrics (percentages, figures, counts)?
   - Does it use active verbs (led, developed, increased) rather than passive voice?

4. COMPLETENESS & CONSISTENCY (weight 15%)
   - Is date formatting consistent throughout the document?
   - Is contact information (email, phone, city, LinkedIn) complete and easy to locate?

5. PROFESSIONAL VISUAL IMPRESSION (weight 10%)
   - Are there visual elements (horizontal divider lines, spacing) that make the CV easy for human eyes to scan, without compromising ATS parsing?

OUTPUT INSTRUCTIONS — provide results EXACTLY in this format (use these exact headings):

## 📊 ATS Score: [score]/100

## 📋 Skor Per Kategori
- Format & Struktur: [x]/25
- Kata Kunci & Relevansi: [x]/30
- Kuantifikasi Hasil: [x]/20
- Kelengkapan & Konsistensi: [x]/15
- Kesan Profesional Visual: [x]/10

## ✅ Kekuatan Utama (3-5 poin)
- [specific finding, NOT generic praise]

## ⚠️ Area yang Perlu Diperbaiki (3-5 poin)
- [concrete and actionable suggestion, NOT vague advice]

## 🔑 Keywords yang Terdeteksi
[list detected keywords/skills found in the CV]

## 📝 Ringkasan Kesiapan ATS
[One paragraph maximum 3 sentences summarizing how ready this CV is to pass ATS filters. No generic praise.]

Respond in Bahasa Indonesia. Focus on SPECIFIC, actionable findings only."""


ATS_CV_PROMPT = """Kamu adalah CV Writer Expert. Berdasarkan CV asli user berikut, buat versi CV yang ATS-friendly.

Rules:
1. Gunakan format yang clean dan terstruktur
2. Gunakan bullet points
3. Highlight skills dan achievements
4. Gunakan keywords yang relevan untuk ATS systems
5. Format sections: PROFIL, PENGALAMAN KERJA, PENDIDIKAN, SKILLS, SERTIFIKASI (jika ada)
6. Tulis dalam bahasa yang sama dengan CV asli (Indonesia/Inggris)

Output HANYA isi CV yang sudah diperbaiki, tanpa penjelasan tambahan.
Gunakan format plain text dengan heading yang jelas."""


def review_cv(cv_text: str, target_job: dict = None) -> dict:
    """
    Analyze CV and return structured feedback.

    Returns dict with:
    - "feedback": AI-generated feedback markdown
    - "available": whether OpenAI/N8N is configured
    """
    # Try N8N first
    if config.is_n8n_configured():
        try:
            from n8n_client import review_cv_n8n
            ai_text = review_cv_n8n(cv_text, target_job=target_job)
            if ai_text and not ai_text.startswith("Error") and not ai_text.startswith("Tidak dapat") and not ai_text.startswith("N8N"):
                return {"feedback": ai_text, "available": True}
        except Exception:
            pass

    if not config.is_llm_configured():
        return {
            "feedback": None,
            "available": False,
        }

    try:
        system_prompt = REVIEW_PROMPT
        target_context = ""
        if target_job:
            target_context = f"\n\nPosisi Target:\n- Jabatan: {target_job.get('job_title', 'N/A')}\n- Perusahaan: {target_job.get('company_name', 'N/A')}\n- Deskripsi Pekerjaan: {target_job.get('job_description', 'N/A')}"
            system_prompt += f"""\n\nIMPORTAN: User menargetkan posisi spesifik. 
Berikan feedback CV yang SPESIFIK dan TERARAH untuk posisi \"{target_job.get('job_title', '')}\". 
Analisis apakah CV user sudah cocok untuk posisi tersebut, identifikasi gap yang perlu diperbaiki, dan berikan saran konkret agar CV lebih menarik untuk posisi ini."""

        reply = chat_completion(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Berikut adalah CV yang perlu di-review:\n\n{cv_text}{target_context}"},
            ],
            temperature=0.5,
            max_tokens=2500,
        )
        return {
            "feedback": reply,
            "available": True,
        }
    except Exception as e:
        return {
            "feedback": f"Error: {str(e)}",
            "available": True,
        }


def generate_ats_cv(cv_text: str, target_job: dict = None, language: str = "auto") -> dict:
    """
    Generate an ATS-friendly version of the CV.

    Returns dict with:
    - "ats_text": improved CV content as plain text
    - "available": whether OpenAI/N8N is configured
    """
    # Try N8N first
    if config.is_n8n_configured():
        try:
            from n8n_client import generate_ats_cv_n8n
            ai_text = generate_ats_cv_n8n(cv_text, target_job=target_job)
            if ai_text and not ai_text.startswith("Error") and not ai_text.startswith("Tidak dapat") and not ai_text.startswith("N8N"):
                return {"ats_text": ai_text, "available": True}
        except Exception:
            pass

    if not config.is_llm_configured():
        return {"ats_text": None, "available": False}

    try:
        system_prompt = ATS_CV_PROMPT
        target_context = ""
        if target_job:
            target_context = f"\n\nPosisi Target:\n- Jabatan: {target_job.get('job_title', 'N/A')}\n- Perusahaan: {target_job.get('company_name', 'N/A')}\n- Deskripsi Pekerjaan: {target_job.get('job_description', 'N/A')}"
            system_prompt += f"""\n\nIMPORTAN: Optimalkan CV ini KHUSUS untuk posisi \"{target_job.get('job_title', '')}\" di \"{target_job.get('company_name', '')}\".
Sesuaikan keywords, skills, dan pengalaman yang di-highlight agar relevan dengan posisi tersebut."""

        if language == "id":
            system_prompt += "\n\nOutput CV ATS-friendly HARUS ditulis dalam Bahasa Indonesia."
        elif language == "en":
            system_prompt += "\n\nOutput CV ATS-friendly HARUS ditulis dalam Bahasa Inggris (English)."

        reply = chat_completion(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"CV Asli:\n\n{cv_text}{target_context}"},
            ],
            temperature=0.4,
            max_tokens=3000,
        )
        return {
            "ats_text": reply,
            "available": True,
        }
    except Exception as e:
        return {"ats_text": f"Error: {str(e)}", "available": True}


def export_cv_to_docx(cv_text: str) -> bytes:
    """Export CV text to a formatted DOCX file. Returns bytes."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)

    # Parse and add content
    lines = cv_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Detect headings (lines in ALL CAPS or starting with ##)
        if line.startswith("## ") or line.startswith("# "):
            clean = line.lstrip("#").strip()
            p = doc.add_heading(clean, level=2)
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            p = doc.add_heading(line, level=2)
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            p = doc.add_paragraph(clean, style="List Bullet")
        else:
            p = doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_cv_to_pdf(cv_text: str) -> bytes:
    """Export CV text to a formatted PDF file. Returns bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    heading_style = ParagraphStyle(
        "CVHeading",
        parent=styles["Heading2"],
        fontSize=13,
        spaceAfter=6,
        spaceBefore=14,
        textColor="#1a1a2e",
    )
    body_style = ParagraphStyle(
        "CVBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=4,
        leading=14,
    )
    bullet_style = ParagraphStyle(
        "CVBullet",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=3,
        leftIndent=20,
        leading=14,
        bulletIndent=10,
    )

    elements = []
    lines = cv_text.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 6))
            continue

        # Escape XML special characters for reportlab
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if line.startswith("## ") or line.startswith("# "):
            clean = safe_line.lstrip("#").strip()
            elements.append(Paragraph(clean, heading_style))
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            elements.append(Paragraph(safe_line, heading_style))
        elif line.startswith("- ") or line.startswith("• "):
            clean = safe_line.lstrip("-•").strip()
            elements.append(Paragraph(f"• {clean}", bullet_style))
        else:
            elements.append(Paragraph(safe_line, body_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
