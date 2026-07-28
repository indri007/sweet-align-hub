"""
CV Processor — Extract text from PDF and DOCX files.
"""

import io
from pathlib import Path
import pdfplumber
from docx import Document
import config


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file.

    Priority:
      1. Gemini Multimodal (Vision) — handles two-column, graphic-heavy CVs natively.
      2. pdfplumber — fast fallback for simple, text-based PDFs.
      3. Tesseract OCR — last resort for fully scanned / image-only PDFs.
    """
    # ── PRIMARY: Gemini Vision ──────────────────────────────────────────────────
    if config.is_gemini_configured():
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=config.GEMINI_API_KEY)
            response = client.models.generate_content(
                model=config.GEMINI_MODEL,
                contents=[
                    types.Part.from_bytes(data=file_bytes, mime_type="application/pdf"),
                    (
                        "Ekstrak SELURUH teks dari dokumen CV ini secara lengkap dan terstruktur. "
                        "Urutkan teks dari atas ke bawah, kiri ke kanan. "
                        "Untuk layout dua kolom, baca kolom kiri terlebih dahulu lalu kolom kanan. "
                        "Pertahankan semua isi teks asli termasuk nama, kontak, pengalaman, pendidikan, "
                        "dan keahlian. Abaikan elemen grafis, ikon, foto profil, dan garis dekoratif. "
                        "Jangan tambahkan komentar atau penjelasan — hanya teks CV saja."
                    ),
                ]
            )
            if response.text and len(response.text.strip()) > 50:
                return response.text.strip()
        except Exception:
            pass  # Gemini rate-limited or unavailable — fall through to pdfplumber

    # ── FALLBACK 1: pdfplumber ──────────────────────────────────────────────────
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=3)
                if page_text:
                    text_parts.append(page_text)
        text = "\n\n".join(text_parts).strip()
        if len(text) >= 50:
            return text
    except Exception:
        pass

    # ── FALLBACK 2: Tesseract OCR (scanned / image-only PDF) ───────────────────
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(file_bytes)
        ocr_parts = [pytesseract.image_to_string(img) for img in images]
        ocr_result = "\n\n".join(ocr_parts).strip()
        if ocr_result:
            return ocr_result
    except Exception:
        pass

    return "\n\n".join(text_parts).strip()



def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_cv_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from a CV file based on its extension.
    Supports PDF and DOCX formats.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".doc":
        raise ValueError(
            "Format .doc (Word lama/binary) tidak didukung. "
            "Silakan simpan ulang CV kamu sebagai .docx atau PDF, lalu upload lagi."
        )
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")


def get_file_info(file_bytes: bytes, filename: str) -> dict:
    """Get basic file info."""
    ext = Path(filename).suffix.lower()
    size_mb = len(file_bytes) / (1024 * 1024)
    info = {
        "filename": filename,
        "format": ext.upper().replace(".", ""),
        "size_mb": round(size_mb, 2),
        "size_bytes": len(file_bytes),
    }
    # Count pages for PDF
    if ext == ".pdf":
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                info["pages"] = len(pdf.pages)
        except Exception:
            info["pages"] = "Unknown"
    elif ext == ".docx":
        try:
            doc = Document(io.BytesIO(file_bytes))
            info["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
        except Exception:
            info["paragraphs"] = "Unknown"
    return info


def validate_cv_file(file_bytes: bytes, filename: str, max_size_mb: int = 100) -> tuple[bool, str]:
    """
    Validate CV file format and size.
    Returns (is_valid, error_message).
    """
    ext = Path(filename).suffix.lower()
    if ext == ".doc":
        return False, "Format .doc (Word lama) tidak didukung. Simpan ulang sebagai .docx atau PDF."
    if ext not in (".pdf", ".docx"):
        return False, f"Format {ext} tidak didukung. Gunakan PDF atau DOCX."

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return False, f"Ukuran file {size_mb:.1f}MB melebihi batas {max_size_mb}MB."

    return True, ""
